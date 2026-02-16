"""DOCX document loader."""

import asyncio

from docx import Document as DocxDocument

from app.rag.loaders.base import BaseDocumentLoader, LoadedDocument


class DocxLoader(BaseDocumentLoader):
    """Loads DOCX documents and extracts text content."""

    SUPPORTED_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def supports_mime_type(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    async def load(self, file_path: str) -> LoadedDocument:
        """Extract text from DOCX."""
        def _extract():
            doc = DocxDocument(file_path)
            paragraphs = []
            tables_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading structure
                    if para.style and para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "")
                        prefix = "#" * int(level) if level.isdigit() else "##"
                        paragraphs.append(f"{prefix} {para.text}")
                    else:
                        paragraphs.append(para.text)

            # Extract table content
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    tables_text.append(f"[Table {i + 1}]\n" + "\n".join(rows))

            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)

            return LoadedDocument(
                text=full_text,
                metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables)},
            )

        return await asyncio.to_thread(_extract)
